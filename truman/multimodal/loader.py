"""
loader.py — Fetch file bytes from DB and build NIM-compatible content blocks.

Return shape from load_attachment():
  {
    "blocks":       list of NIM content blocks (image_url or text),
    "text_inline":  str — extracted text for PDF/DOCX/sheet/code (empty for pure images),
    "tokens_est":   int — rough token estimate,
    "kind":         str — "image" | "pdf_text" | "pdf_scan" | "docx" | "xlsx" | "csv"
                           | "code" | "text" | "unknown",
  }
"""
from __future__ import annotations
import base64
import io

# ── Constants ─────────────────────────────────────────────────────────────────
MAX_PDF_PAGES      = 20       # cap for scanned PDFs sent as images
MAX_SHEET_ROWS     = 200      # cap for XLSX/CSV → markdown table
MAX_TEXT_CHARS     = 12_000   # cap for raw text/code inline


# ── Main public function ──────────────────────────────────────────────────────

def load_attachment(attach_id: str, page_hint: int | None = None) -> dict:
    """
    Full type-aware loader. Returns the unified dict described in module docstring.
    page_hint: if set, for PDFs only return that 1-indexed page.
    """
    try:
        from truman.storage.db import get_attachment
        att = get_attachment(attach_id)
        if not att:
            print(f"[Multimodal] attach_id {attach_id} not found in DB")
            return _empty("unknown")

        mime     = att.get("mime_type", "application/octet-stream")
        filename = att.get("filename", "")
        data: bytes = att["data"]

        # ── Route by mime / extension ─────────────────────────────────────────
        if mime.startswith("image/"):
            return _load_image(data, mime)

        if mime == "application/pdf" or filename.lower().endswith(".pdf"):
            return _load_pdf(data, page_hint=page_hint)

        if mime in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "application/msword") or filename.lower().endswith((".docx", ".doc")):
            return _load_docx(data)

        if mime in ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    "application/vnd.ms-excel") or filename.lower().endswith((".xlsx", ".xls")):
            return _load_xlsx(data)

        if mime == "text/csv" or filename.lower().endswith(".csv"):
            return _load_csv(data)

        if _is_code_file(filename) or mime.startswith("text/"):
            return _load_text(data, filename)

        # Unknown binary — return as text best-effort
        return _load_text(data, filename)

    except Exception as e:
        print(f"[Multimodal] loader error for {attach_id}: {e}")
        return _empty("unknown")


# ── Legacy compat: single image block (used by Phase 1 nodes.py wiring) ──────

def load_image_block(attach_id: str) -> dict | None:
    """
    Legacy function — returns a single NIM image_url block or None.
    Call load_attachment() for the full type-aware result.
    """
    try:
        from truman.storage.db import get_attachment
        att = get_attachment(attach_id)
        if not att:
            return None
        mime = att["mime_type"]
        if not mime.startswith("image/"):
            return None
        b64 = base64.b64encode(att["data"]).decode()
        return {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}}
    except Exception as e:
        print(f"[Multimodal] load_image_block error for {attach_id}: {e}")
        return None


def get_attachment_meta(attach_id: str) -> dict | None:
    """Return filename + mime for an attach_id (no bytes)."""
    try:
        from truman.storage.db import get_attachment
        att = get_attachment(attach_id)
        if not att:
            return None
        return {"filename": att.get("filename", ""), "mime": att.get("mime_type", "")}
    except Exception:
        return None


# ── Type-specific loaders ─────────────────────────────────────────────────────

def _load_image(data: bytes, mime: str) -> dict:
    b64 = base64.b64encode(data).decode()
    block = {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}}
    return {"blocks": [block], "text_inline": "", "tokens_est": 768, "kind": "image"}


def _load_pdf(data: bytes, page_hint: int | None = None) -> dict:
    """
    Try text extraction first (pdfplumber). If extracted text is too sparse
    (< 50 chars per page avg), treat as scanned and send pages as images.
    page_hint: 1-indexed. If set, only that page is processed.
    """
    import pdfplumber

    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            total_pages = len(pdf.pages)

            if page_hint is not None:
                # Pin to a single page
                idx = min(max(page_hint - 1, 0), total_pages - 1)
                pages_to_use = [pdf.pages[idx]]
                page_range_label = f"(page {page_hint})"
            else:
                pages_to_use = pdf.pages[:MAX_PDF_PAGES]
                page_range_label = ""

            # Attempt text extraction
            all_text = []
            for p in pages_to_use:
                t = p.extract_text() or ""
                all_text.append(t)

            combined = "\n\n".join(all_text).strip()
            avg_chars = len(combined) / max(len(pages_to_use), 1)

            if avg_chars >= 50:
                # Good text PDF
                text = combined[:MAX_TEXT_CHARS]
                if len(combined) > MAX_TEXT_CHARS:
                    text += f"\n\n[...truncated — {total_pages} pages total{page_range_label}]"
                label = f"[PDF — {total_pages} pages{page_range_label}]\n\n" + text
                return {
                    "blocks":      [{"type": "text", "text": label}],
                    "text_inline": label,
                    "tokens_est":  len(label) // 4,
                    "kind":        "pdf_text",
                }

            # Scanned PDF — render pages as images
            return _pdf_as_images(data, pages_to_use, total_pages, page_range_label)

    except Exception as e:
        print(f"[Multimodal] pdfplumber error: {e}")
        return _empty("pdf_text")


def _pdf_as_images(data: bytes, pages_to_render, total_pages: int, label: str) -> dict:
    """Convert PDF pages to PNG and return as image_url blocks."""
    try:
        import fitz  # PyMuPDF — optional dep
        doc = fitz.open(stream=data, filetype="pdf")
        blocks = []
        for page in pages_to_render:
            pno = page.number
            mat = fitz.Matrix(1.5, 1.5)   # 1.5x zoom → readable without being huge
            pix = doc[pno].get_pixmap(matrix=mat)
            png_bytes = pix.tobytes("png")
            b64 = base64.b64encode(png_bytes).decode()
            blocks.append({"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}})
        intro = {"type": "text", "text": f"[Scanned PDF — {total_pages} pages{label}. Showing {len(blocks)} page(s) as images]"}
        return {
            "blocks":      [intro] + blocks,
            "text_inline": "",
            "tokens_est":  768 * len(blocks),
            "kind":        "pdf_scan",
        }
    except ImportError:
        # PyMuPDF not installed — return the extracted (sparse) text anyway
        return _empty("pdf_scan")
    except Exception as e:
        print(f"[Multimodal] PDF→image error: {e}")
        return _empty("pdf_scan")


def _load_docx(data: bytes) -> dict:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                style = para.style.name.lower()
                if "heading" in style:
                    level = style.replace("heading", "").strip()
                    prefix = "#" * (int(level) if level.isdigit() else 1)
                    lines.append(f"{prefix} {para.text.strip()}")
                else:
                    lines.append(para.text.strip())
        # Tables
        for table in doc.tables:
            rows = [[c.text.strip() for c in row.cells] for row in table.rows]
            if rows:
                header = "| " + " | ".join(rows[0]) + " |"
                sep    = "| " + " | ".join(["---"] * len(rows[0])) + " |"
                body   = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
                lines += [header, sep, body]

        text = "\n".join(lines)[:MAX_TEXT_CHARS]
        label = "[DOCX]\n\n" + text
        return {
            "blocks":      [{"type": "text", "text": label}],
            "text_inline": label,
            "tokens_est":  len(label) // 4,
            "kind":        "docx",
        }
    except Exception as e:
        print(f"[Multimodal] docx error: {e}")
        return _empty("docx")


def _load_xlsx(data: bytes) -> dict:
    try:
        import pandas as pd
        xl = pd.ExcelFile(io.BytesIO(data))
        parts = []
        for sheet_name in xl.sheet_names[:5]:  # max 5 sheets
            df = xl.parse(sheet_name).head(MAX_SHEET_ROWS)
            md = df.to_markdown(index=False)
            parts.append(f"### Sheet: {sheet_name}\n\n{md}")
        text = "\n\n".join(parts)[:MAX_TEXT_CHARS]
        label = "[XLSX]\n\n" + text
        return {
            "blocks":      [{"type": "text", "text": label}],
            "text_inline": label,
            "tokens_est":  len(label) // 4,
            "kind":        "xlsx",
        }
    except Exception as e:
        print(f"[Multimodal] xlsx error: {e}")
        return _empty("xlsx")


def _load_csv(data: bytes) -> dict:
    try:
        import pandas as pd
        df = pd.read_csv(io.BytesIO(data)).head(MAX_SHEET_ROWS)
        md = df.to_markdown(index=False)
        text = ("[CSV]\n\n" + md)[:MAX_TEXT_CHARS]
        return {
            "blocks":      [{"type": "text", "text": text}],
            "text_inline": text,
            "tokens_est":  len(text) // 4,
            "kind":        "csv",
        }
    except Exception as e:
        print(f"[Multimodal] csv error: {e}")
        return _empty("csv")


def _load_text(data: bytes, filename: str) -> dict:
    try:
        text = data.decode("utf-8", errors="replace")[:MAX_TEXT_CHARS]
        ext  = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        kind = "code" if _is_code_file(filename) else "text"
        lang = ext if kind == "code" else ""
        fence = f"```{lang}\n{text}\n```" if lang else text
        label = f"[{kind.upper()} — {filename}]\n\n{fence}"
        return {
            "blocks":      [{"type": "text", "text": label}],
            "text_inline": label,
            "tokens_est":  len(label) // 4,
            "kind":        kind,
        }
    except Exception as e:
        print(f"[Multimodal] text load error: {e}")
        return _empty("text")


# ── Helpers ───────────────────────────────────────────────────────────────────

_CODE_EXTS = {
    "py", "js", "ts", "tsx", "jsx", "go", "rs", "java", "kt", "swift",
    "c", "cpp", "h", "hpp", "cs", "rb", "php", "sh", "bash", "zsh",
    "sql", "html", "css", "scss", "json", "yaml", "yml", "toml", "env",
    "dockerfile", "makefile",
}

def _is_code_file(filename: str) -> bool:
    if not filename:
        return False
    lower = filename.lower()
    if "." not in lower:
        return lower in ("dockerfile", "makefile", "procfile")
    ext = lower.rsplit(".", 1)[-1]
    return ext in _CODE_EXTS


def _empty(kind: str) -> dict:
    return {"blocks": [], "text_inline": "", "tokens_est": 0, "kind": kind}
